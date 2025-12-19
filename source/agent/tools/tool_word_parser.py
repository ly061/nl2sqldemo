"""
Word文档解析工具
用于解析用户上传的Word文档，提取文本内容供测试用例生成使用
"""
import base64
import tempfile
from pathlib import Path
from typing import Optional, Tuple, List
from docx import Document
from langchain_core.tools import tool
from source.agent.utils.log_utils import MyLogger

log = MyLogger().get_logger()


def _parse_word_from_path(doc_path: Path) -> Tuple[List[str], List[str]]:
    """从文件路径解析Word文档
    
    Returns:
        (paragraphs, tables_content) 元组
    """
    doc = Document(doc_path)
    
    # 提取所有段落文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # 忽略空段落
            paragraphs.append(text)
    
    # 提取表格内容
    tables_content = []
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            if any(row_cells):  # 忽略空行
                table_rows.append(" | ".join(row_cells))
        if table_rows:
            tables_content.append("\n".join(table_rows))
    
    return paragraphs, tables_content


@tool
def parse_word_document(file_path: str, base64_data: Optional[str] = None) -> str:
    """解析Word文档，提取文本内容
    
    支持两种方式：
    1. 通过文件路径解析（file_path参数）
    2. 通过base64数据解析（base64_data参数，会自动创建临时文件）
    
    Args:
        file_path: Word文档的文件路径（如果提供base64_data，此参数可为空）
        base64_data: Word文档的base64编码数据（可选）
    
    Returns:
        提取的文本内容
    """
    try:
        temp_file = None
        
        # 如果提供了base64数据，先保存为临时文件
        if base64_data:
            try:
                # 解码base64数据
                file_data = base64.b64decode(base64_data)
                
                # 创建临时文件
                temp_file = tempfile.NamedTemporaryFile(
                    delete=False,
                    suffix='.docx',
                    prefix='word_'
                )
                temp_file.write(file_data)
                temp_file.close()
                
                doc_path = Path(temp_file.name)
                log.info(f"从base64数据创建临时Word文件: {doc_path}")
            except Exception as e:
                if temp_file:
                    try:
                        Path(temp_file.name).unlink(missing_ok=True)
                    except:
                        pass
                return f"错误：base64数据解码失败 - {str(e)}"
        else:
            if not file_path:
                return "错误：必须提供file_path或base64_data参数"
            
            doc_path = Path(file_path)
            if not doc_path.exists():
                return f"错误：文件不存在 - {file_path}"
        
        # 检查文件格式
        if doc_path.suffix.lower() not in ['.docx', '.doc']:
            if temp_file:
                try:
                    Path(temp_file.name).unlink(missing_ok=True)
                except:
                    pass
            return f"错误：不支持的文件格式 - {doc_path.suffix}，仅支持 .docx 或 .doc 格式"
        
        # 解析Word文档
        paragraphs, tables_content = _parse_word_from_path(doc_path)
        
        # 清理临时文件
        if temp_file:
            try:
                Path(temp_file.name).unlink(missing_ok=True)
            except Exception as e:
                log.warning(f"清理临时文件失败: {e}")
        
        # 组合所有内容
        content_parts = []
        if paragraphs:
            content_parts.append("\n".join(paragraphs))
        if tables_content:
            content_parts.append("\n\n表格内容：\n" + "\n\n".join(tables_content))
        
        full_content = "\n\n".join(content_parts) if content_parts else "文档为空"
        
        log.info(f"成功解析Word文档: {doc_path}，提取了 {len(paragraphs)} 个段落和 {len(tables_content)} 个表格")
        
        return f"Word文档内容提取成功：\n\n{full_content}"
    
    except Exception as e:
        error_msg = f"解析Word文档失败: {str(e)}"
        log.error(error_msg)
        # 确保清理临时文件
        if temp_file:
            try:
                Path(temp_file.name).unlink(missing_ok=True)
            except:
                pass
        return error_msg

