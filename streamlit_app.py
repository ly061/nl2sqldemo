import asyncio
import os
import urllib.parse
import uuid
import base64
import tempfile
from collections.abc import AsyncGenerator
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv
from pydantic import ValidationError

# 尝试导入 docx 库用于解析 Word 文档
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from agent_client import AgentClient, AgentClientError
from api.schema import ChatHistory, ChatMessage
try:
    from schema.task_data import TaskData, TaskDataStatus
except ImportError:
    # 占位符：如果 task_data 模块不存在，创建简单的占位符
    from pydantic import BaseModel
    from typing import Any, Dict
    
    class TaskData(BaseModel):
        custom_data: Dict[str, Any]
    
    class TaskDataStatus:
        def __init__(self):
            pass
        def add_and_draw_task_data(self, task_data: TaskData):
            import streamlit as st
            st.write("Task data:", task_data.custom_data)

try:
    from voice import VoiceManager
except ImportError:
    # 占位符：如果 voice 模块不存在，创建简单的占位符
    class VoiceManager:
        @staticmethod
        def from_env():
            return None
        
        def get_chat_input(self):
            return None
        
        def render_message(self, content, container=None, audio_only=False):
            if container:
                container.write(content)
            else:
                st.write(content)

# A Streamlit app for interacting with the langgraph agent via a simple chat interface.
# The app has three main functions which are all run async:

# - main() - sets up the streamlit app and high level structure
# - draw_messages() - draws a set of chat messages - either replaying existing messages
#   or streaming new ones.
# - handle_feedback() - Draws a feedback widget and records feedback from the user.

# The app heavily uses AgentClient to interact with the agent's FastAPI endpoints.


APP_TITLE = "Agent Service Toolkit"
APP_ICON = "🧰"
USER_ID_COOKIE = "user_id"


def parse_word_document(file_bytes: bytes) -> str:
    """解析 Word 文档内容
    
    Args:
        file_bytes: Word 文档的字节数据
    
    Returns:
        提取的文本内容
    """
    if not DOCX_AVAILABLE:
        return "错误：python-docx 库未安装"
    
    try:
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(file_bytes)
        temp_file.close()
        
        try:
            # 解析 Word 文档
            doc = Document(temp_file.name)
            
            # 提取所有段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 忽略空段落
                    paragraphs.append(text)
            
            # 提取表格内容
            tables_content = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):  # 忽略空行
                        table_rows.append(" | ".join(row_cells))
                if table_rows:
                    tables_content.append("\n".join(table_rows))
            
            # 组合所有内容
            content_parts = []
            if paragraphs:
                content_parts.append("\n".join(paragraphs))
            if tables_content:
                content_parts.append("\n\n表格内容：\n" + "\n\n".join(tables_content))
            
            full_content = "\n\n".join(content_parts) if content_parts else "文档为空"
            
            return full_content
        finally:
            # 清理临时文件
            try:
                Path(temp_file.name).unlink(missing_ok=True)
            except:
                pass
    except Exception as e:
        return f"解析 Word 文档失败: {str(e)}"


def get_or_create_user_id() -> str:
    """Get the user ID from session state or URL parameters, or create a new one if it doesn't exist."""
    # Check if user_id exists in session state
    if USER_ID_COOKIE in st.session_state:
        return st.session_state[USER_ID_COOKIE]

    # Try to get from URL parameters using the new st.query_params
    if USER_ID_COOKIE in st.query_params:
        user_id = st.query_params[USER_ID_COOKIE]
        st.session_state[USER_ID_COOKIE] = user_id
        return user_id

    # Generate a new user_id if not found
    user_id = str(uuid.uuid4())

    # Store in session state for this session
    st.session_state[USER_ID_COOKIE] = user_id

    # Also add to URL parameters so it can be bookmarked/shared
    st.query_params[USER_ID_COOKIE] = user_id

    return user_id


async def main() -> None:
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon=APP_ICON,
        menu_items={},
    )

    # Hide the streamlit upper-right chrome and add thinking animation
    st.html(
        """
        <style>
        [data-testid="stStatusWidget"] {
                visibility: hidden;
                height: 0%;
                position: fixed;
            }
        
        /* Thinking animation styles */
        .thinking-container {
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 8px 0;
        }
        
        .thinking-dots {
            display: flex;
            gap: 4px;
        }
        
        .thinking-dots span {
            width: 8px;
            height: 8px;
            background-color: #6366f1;
            border-radius: 50%;
            animation: bounce 1.4s infinite ease-in-out both;
        }
        
        .thinking-dots span:nth-child(1) {
            animation-delay: -0.32s;
        }
        
        .thinking-dots span:nth-child(2) {
            animation-delay: -0.16s;
        }
        
        .thinking-dots span:nth-child(3) {
            animation-delay: 0s;
        }
        
        @keyframes bounce {
            0%, 80%, 100% {
                transform: scale(0);
            }
            40% {
                transform: scale(1);
            }
        }
        
        .thinking-text {
            color: #6366f1;
            font-weight: 500;
            animation: pulse 2s infinite;
        }
        
        @keyframes pulse {
            0%, 100% {
                opacity: 1;
            }
            50% {
                opacity: 0.5;
            }
        }
        </style>
        """,
    )
    if st.get_option("client.toolbarMode") != "minimal":
        st.set_option("client.toolbarMode", "minimal")
        await asyncio.sleep(0.1)
        st.rerun()

    # Get or create user ID
    user_id = get_or_create_user_id()

    if "agent_client" not in st.session_state:
        load_dotenv()
        agent_url = os.getenv("AGENT_URL")
        if not agent_url:
            # 使用 localhost 而不是 0.0.0.0（0.0.0.0 是服务器绑定地址，不能用于客户端连接）
            host = os.getenv("HOST", "localhost")
            port = os.getenv("PORT", 9501)  # 修复：默认端口应该是 9501，不是 9051
            agent_url = f"http://{host}:{port}"
        try:
            with st.spinner("Connecting to agent service..."):
                st.session_state.agent_client = AgentClient(base_url=agent_url)
        except AgentClientError as e:
            st.error(f"Error connecting to agent service at {agent_url}: {e}")
            st.markdown("The service might be booting up. Try again in a few seconds.")
            st.stop()
    agent_client: AgentClient = st.session_state.agent_client

    # Initialize voice manager (once per session)
    if "voice_manager" not in st.session_state:
        st.session_state.voice_manager = VoiceManager.from_env()
    voice = st.session_state.voice_manager

    if "thread_id" not in st.session_state:
        thread_id = st.query_params.get("thread_id")
        if not thread_id:
            thread_id = str(uuid.uuid4())
            messages = []
        else:
            try:
                messages: ChatHistory = agent_client.get_history(thread_id=thread_id).messages
            except AgentClientError:
                st.error("No message history found for this Thread ID.")
                messages = []
        st.session_state.messages = messages
        st.session_state.thread_id = thread_id

    # 默认配置
    use_streaming = True
    enable_audio = False
    
    # 简化侧边栏：只保留 New Chat 按钮
    with st.sidebar:
        st.header(f"{APP_ICON} {APP_TITLE}")

        ""
        "测试用例生成系统"
        ""

        if st.button(":material/chat: New Chat", use_container_width=True):
            st.session_state.messages = []
            st.session_state.thread_id = str(uuid.uuid4())
            # Clear saved audio when starting new chat
            if "last_audio" in st.session_state:
                del st.session_state.last_audio
            st.rerun()

    # Draw existing messages
    messages: list[ChatMessage] = st.session_state.messages

    if len(messages) == 0:
        # 简化欢迎消息（当前系统是测试用例生成系统）
        WELCOME = "Hello! I'm a test case generation assistant. I can help you generate test cases from requirements. Ask me anything!"

        with st.chat_message("ai"):
            st.write(WELCOME)

    # draw_messages() expects an async iterator over messages
    async def amessage_iter() -> AsyncGenerator[ChatMessage, None]:
        for m in messages:
            yield m

    await draw_messages(amessage_iter())

    # Render saved audio for the last AI message (if it exists)
    # This ensures audio persists across st.rerun() calls
    if (
        voice
        and enable_audio
        and "last_audio" in st.session_state
        and st.session_state.last_message
        and len(messages) > 0
        and messages[-1].type == "ai"
    ):
        with st.session_state.last_message:
            audio_data = st.session_state.last_audio
            st.audio(audio_data["data"], format=audio_data["format"])

    # Generate new message if the user provided new input
    # Use voice manager if available, otherwise fall back to regular input
    # REQUIRED: Set VOICE_STT_PROVIDER, VOICE_TTS_PROVIDER, OPENAI_API_KEY
    # in app .env (NOT service .env) to enable voice features.
    
    # 初始化 file_uploader 的动态 key（用于重置上传组件）
    if "file_uploader_key" not in st.session_state:
        st.session_state.file_uploader_key = 0
    
    # 创建并排布局：输入框和可折叠的文件上传在同一行
    if voice:
        user_input = voice.get_chat_input()
        uploaded_file = None
    else:
        # 使用列布局：输入框在左侧，文件上传在右侧
        col1, col2 = st.columns([0.9, 0.1])
        
        with col1:
            user_input = st.chat_input(placeholder="输入需求或上传 Word 文档...")
        
        with col2:
            # 使用 popover 创建可折叠的文件上传
            with st.popover("📄", use_container_width=True, help="上传 Word 文档"):
                # 使用动态 key，这样可以通过改变 key 来重置上传组件
                uploaded_file = st.file_uploader(
                    "上传 Word 需求文档",
                    type=['docx'],
                    help="支持上传 Word 文档，系统会自动解析文档内容并用于生成测试用例",
                    key=f"word_file_uploader_{st.session_state.file_uploader_key}"
                )
                if uploaded_file is not None:
                    st.success(f"✅ 已上传: {uploaded_file.name}")
    
    # 处理文件上传
    if uploaded_file is not None:
        # 将文件内容保存到 session state
        if "uploaded_file_content" not in st.session_state or st.session_state.get("uploaded_file_name") != uploaded_file.name:
            st.session_state.uploaded_file_content = uploaded_file.read()
            st.session_state.uploaded_file_name = uploaded_file.name
            st.session_state.file_parsed = False

    # 处理文件上传和用户输入
    processed_input = None
    if user_input:  # 只有当用户有输入时才处理
        if "uploaded_file_content" in st.session_state and st.session_state.uploaded_file_content:
            # 有文件上传，解析文档
            try:
                if DOCX_AVAILABLE:
                    # 解析 Word 文档
                    doc_content = parse_word_document(st.session_state.uploaded_file_content)
                    file_name = st.session_state.get("uploaded_file_name", "文档")
                    
                    # 合并用户输入和文档内容
                    processed_input = f"需求文档《{file_name}》内容：\n{doc_content}\n\n用户补充说明：\n{user_input}"
                    st.info(f"📄 已解析文档《{file_name}》，内容已添加到输入中")
                else:
                    st.warning("⚠️ python-docx 库未安装，无法解析 Word 文档。请安装: pip install python-docx")
                    processed_input = user_input
            except Exception as e:
                st.error(f"❌ 解析 Word 文档失败: {str(e)}")
                processed_input = user_input
        else:
            # 没有文件上传，直接使用用户输入
            processed_input = user_input
    elif "uploaded_file_content" in st.session_state and st.session_state.uploaded_file_content:
        # 只有文件上传，没有文本输入，提示用户输入
        st.info("📄 已上传文档，请在输入框中输入需求或直接发送消息以基于文档生成测试用例")

    if processed_input:
        messages.append(ChatMessage(type="human", content=processed_input))
        st.chat_message("human").write(processed_input)
        # 清除文件状态，确保下次消息不再携带文档内容
        # 同时增加 file_uploader_key 来重置上传组件
        if "uploaded_file_content" in st.session_state:
            del st.session_state.uploaded_file_content
            # 增加 key 来重置 file_uploader 组件
            st.session_state.file_uploader_key = st.session_state.get("file_uploader_key", 0) + 1
        if "uploaded_file_name" in st.session_state:
            del st.session_state.uploaded_file_name
        if "file_parsed" in st.session_state:
            del st.session_state.file_parsed
        try:
            if use_streaming:
                stream = agent_client.astream(
                    message=processed_input,
                    thread_id=st.session_state.thread_id,
                )
                await draw_messages(stream, is_new=True)
                # Generate TTS audio for streaming response
                # Note: draw_messages() stores the final message in st.session_state.messages
                # and the container reference in st.session_state.last_message
                if voice and enable_audio and st.session_state.messages:
                    last_msg = st.session_state.messages[-1]
                    # Only generate audio for AI responses with content
                    if last_msg.type == "ai" and last_msg.content:
                        # Use audio_only=True since text was already streamed by draw_messages()
                        voice.render_message(
                            last_msg.content,
                            container=st.session_state.last_message,
                            audio_only=True,
                        )
            else:
                response = await agent_client.ainvoke(
                    message=processed_input,
                    thread_id=st.session_state.thread_id,
                )
                messages.append(response)
                # Render AI response with optional voice
                with st.chat_message("ai"):
                    if voice and enable_audio:
                        voice.render_message(response.content)
                    else:
                        st.write(response.content)
            st.rerun()  # Clear stale containers
        except AgentClientError as e:
            st.error(f"Error generating response: {e}")
            st.stop()

    # If messages have been generated, show feedback widget
    if len(messages) > 0 and st.session_state.last_message:
        with st.session_state.last_message:
            await handle_feedback()


async def draw_messages(
    messages_agen: AsyncGenerator[ChatMessage | str, None],
    is_new: bool = False,
) -> None:
    """
    Draws a set of chat messages - either replaying existing messages
    or streaming new ones.

    This function has additional logic to handle streaming tokens and tool calls.
    - Use a placeholder container to render streaming tokens as they arrive.
    - Use a status container to render tool calls. Track the tool inputs and outputs
      and update the status container accordingly.

    The function also needs to track the last message container in session state
    since later messages can draw to the same container. This is also used for
    drawing the feedback widget in the latest chat message.

    Args:
        messages_aiter: An async iterator over messages to draw.
        is_new: Whether the messages are new or not.
    """

    # Keep track of the last message container
    last_message_type = None
    st.session_state.last_message = None

    # Placeholder for intermediate streaming tokens
    streaming_content = ""
    streaming_placeholder = None
    
    # Thinking indicator for new messages
    thinking_placeholder = None
    if is_new:
        # Show thinking indicator while waiting for AI response
        st.session_state.last_message = st.chat_message("ai")
        with st.session_state.last_message:
            thinking_placeholder = st.empty()
            # Use animated thinking indicator
            thinking_placeholder.html("""
                <div class="thinking-container">
                    <div class="thinking-dots">
                        <span></span>
                        <span></span>
                        <span></span>
                    </div>
                    <span class="thinking-text">AI 正在思考...</span>
                </div>
            """)
        last_message_type = "ai"

    # Iterate over the messages and draw them
    while msg := await anext(messages_agen, None):
        # Clear thinking indicator on first message received
        if thinking_placeholder:
            thinking_placeholder.empty()
            thinking_placeholder = None
            
        # str message represents an intermediate token being streamed
        if isinstance(msg, str):
            # If placeholder is empty, this is the first token of a new message
            # being streamed. We need to do setup.
            if not streaming_placeholder:
                if last_message_type != "ai":
                    last_message_type = "ai"
                    st.session_state.last_message = st.chat_message("ai")
                with st.session_state.last_message:
                    streaming_placeholder = st.empty()

            streaming_content += msg
            streaming_placeholder.write(streaming_content)
            continue
        if not isinstance(msg, ChatMessage):
            st.error(f"Unexpected message type: {type(msg)}")
            st.write(msg)
            st.stop()

        match msg.type:
            # A message from the user, the easiest case
            case "human":
                last_message_type = "human"
                st.chat_message("human").write(msg.content)

            # A message from the agent is the most complex case, since we need to
            # handle streaming tokens and tool calls.
            case "ai":
                # If we're rendering new messages, store the message in session state
                if is_new:
                    st.session_state.messages.append(msg)

                # If the last message type was not AI, create a new chat message
                if last_message_type != "ai":
                    last_message_type = "ai"
                    st.session_state.last_message = st.chat_message("ai")

                with st.session_state.last_message:
                    # If the message has content, write it out.
                    # Reset the streaming variables to prepare for the next message.
                    # Debug: Check if content is None or empty
                    if msg.content is None:
                        st.warning("⚠️ 收到空内容 (None)，可能是本地模型响应格式问题")
                        # Log for debugging
                        import logging
                        logging.getLogger(__name__).warning(f"AI message with None content: {msg.model_dump()}")
                    elif msg.content == "":
                        st.info("ℹ️ 收到空字符串内容")
                    elif msg.content:
                        if streaming_placeholder:
                            streaming_placeholder.write(msg.content)
                            streaming_content = ""
                            streaming_placeholder = None
                        else:
                            st.write(msg.content)

                    if msg.tool_calls:
                        # Create a status container for each tool call and store the
                        # status container by ID to ensure results are mapped to the
                        # correct status container.
                        call_results = {}
                        for tool_call in msg.tool_calls:
                            # Use different labels for transfer vs regular tool calls
                            if "transfer_to" in tool_call["name"]:
                                label = f"""💼 Sub Agent: {tool_call["name"]}"""
                            else:
                                label = f"""🛠️ Tool Call: {tool_call["name"]}"""

                            status = st.status(
                                label,
                                state="running" if is_new else "complete",
                                expanded=False,  # 默认折叠
                            )
                            call_results[tool_call["id"]] = status

                        # Expect one ToolMessage for each tool call.
                        for tool_call in msg.tool_calls:
                            if "transfer_to" in tool_call["name"]:
                                status = call_results[tool_call["id"]]
                                status.update(expanded=False)  # 默认折叠
                                await handle_sub_agent_msgs(messages_agen, status, is_new)
                                break

                            # Only non-transfer tool calls reach this point
                            status = call_results[tool_call["id"]]
                            status.write("Input:")
                            status.write(tool_call["args"])
                            tool_result: ChatMessage = await anext(messages_agen)

                            if tool_result.type != "tool":
                                st.error(f"Unexpected ChatMessage type: {tool_result.type}")
                                st.write(tool_result)
                                st.stop()

                            # Record the message if it's new, and update the correct
                            # status container with the result
                            if is_new:
                                st.session_state.messages.append(tool_result)
                            if tool_result.tool_call_id:
                                status = call_results[tool_result.tool_call_id]
                            status.write("Output:")
                            status.write(tool_result.content)
                            status.update(state="complete")

            case "custom":
                # CustomData example used by the bg-task-agent
                # See:
                # - src/agents/utils.py CustomData
                # - src/agents/bg_task_agent/task.py
                try:
                    task_data: TaskData = TaskData.model_validate(msg.custom_data)
                except ValidationError:
                    st.error("Unexpected CustomData message received from agent")
                    st.write(msg.custom_data)
                    st.stop()

                if is_new:
                    st.session_state.messages.append(msg)

                if last_message_type != "task":
                    last_message_type = "task"
                    st.session_state.last_message = st.chat_message(
                        name="task", avatar=":material/manufacturing:"
                    )
                    with st.session_state.last_message:
                        status = TaskDataStatus()

                status.add_and_draw_task_data(task_data)

            # In case of an unexpected message type, log an error and stop
            case _:
                st.error(f"Unexpected ChatMessage type: {msg.type}")
                st.write(msg)
                st.stop()


async def handle_feedback() -> None:
    """Draws a feedback widget and records feedback from the user."""

    # Keep track of last feedback sent to avoid sending duplicates
    if "last_feedback" not in st.session_state:
        st.session_state.last_feedback = (None, None)

    # 使用消息索引作为key（因为ChatMessage没有run_id）
    latest_message_idx = len(st.session_state.messages) - 1
    feedback = st.feedback("stars", key=f"feedback_{latest_message_idx}")

    # If the feedback value has changed, record it
    if feedback is not None and (latest_message_idx, feedback) != st.session_state.last_feedback:
        # Normalize the feedback value (an index) to a score between 0 and 1
        normalized_score = (feedback + 1) / 5.0

        # 简化反馈：只记录到session state，不发送到API（因为API不支持feedback端点）
        st.session_state.last_feedback = (latest_message_idx, feedback)
        st.toast(f"Feedback recorded: {feedback} stars", icon=":material/reviews:")


async def handle_sub_agent_msgs(messages_agen, status, is_new):
    """
    This function segregates agent output into a status container.
    It handles all messages after the initial tool call message
    until it reaches the final AI message.

    Enhanced to support nested multi-agent hierarchies with handoff back messages.

    Args:
        messages_agen: Async generator of messages
        status: the status container for the current agent
        is_new: Whether messages are new or replayed
    """
    nested_popovers = {}

    # looking for the transfer Success tool call message
    first_msg = await anext(messages_agen)
    if is_new:
        st.session_state.messages.append(first_msg)

    # Continue reading until we get an explicit handoff back
    while True:
        # Read next message
        sub_msg = await anext(messages_agen)

        # this should only happen is skip_stream flag is removed
        # if isinstance(sub_msg, str):
        #     continue

        if is_new:
            st.session_state.messages.append(sub_msg)

        # Handle tool results with nested popovers
        if sub_msg.type == "tool" and sub_msg.tool_call_id in nested_popovers:
            popover = nested_popovers[sub_msg.tool_call_id]
            popover.write("**Output:**")
            popover.write(sub_msg.content)
            continue

        # Handle transfer_back_to tool calls - these indicate a sub-agent is returning control
        if (
            hasattr(sub_msg, "tool_calls")
            and sub_msg.tool_calls
            and any("transfer_back_to" in tc.get("name", "") for tc in sub_msg.tool_calls)
        ):
            # Process transfer_back_to tool calls
            for tc in sub_msg.tool_calls:
                if "transfer_back_to" in tc.get("name", ""):
                    # Read the corresponding tool result
                    transfer_result = await anext(messages_agen)
                    if is_new:
                        st.session_state.messages.append(transfer_result)

            # After processing transfer back, we're done with this agent
            if status:
                status.update(state="complete")
            break

        # Display content and tool calls in the same nested status
        if status:
            if sub_msg.content:
                status.write(sub_msg.content)

            if hasattr(sub_msg, "tool_calls") and sub_msg.tool_calls:
                for tc in sub_msg.tool_calls:
                    # Check if this is a nested transfer/delegate
                    if "transfer_to" in tc["name"]:
                        # Create a nested status container for the sub-agent
                        nested_status = status.status(
                            f"""💼 Sub Agent: {tc["name"]}""",
                            state="running" if is_new else "complete",
                            expanded=False,  # 默认折叠
                        )

                        # Recursively handle sub-agents of this sub-agent
                        await handle_sub_agent_msgs(messages_agen, nested_status, is_new)
                    else:
                        # Regular tool call - create popover
                        popover = status.popover(f"{tc['name']}", icon="🛠️")
                        popover.write(f"**Tool:** {tc['name']}")
                        popover.write("**Input:**")
                        popover.write(tc["args"])
                        # Store the popover reference using the tool call ID
                        nested_popovers[tc["id"]] = popover


if __name__ == "__main__":
    asyncio.run(main())